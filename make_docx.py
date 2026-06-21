import sys
import os
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_rtl(p):
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pPr = p._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    pPr.append(bidi)

def make_word_doc(html_path, out_path):
    with open(html_path, 'r', encoding='utf-8') as f:
        html_content = f.read()
    
    soup = BeautifulSoup(html_content, 'html.parser')
    doc_body = soup.find('div', class_='doc-body')
    doc_header = soup.find('div', class_='doc-header')
    
    doc = Document()

    # Title
    if doc_header:
        h1 = doc_header.find('h1')
        if h1:
            title = doc.add_heading(h1.get_text(), level=0)
            set_rtl(title)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in title.runs:
                run.font.name = 'Arial'
                run.font.color.rgb = RGBColor(0, 0, 0)
        
        subtitle = doc_header.find('p', class_='subtitle')
        if subtitle:
            p = doc.add_paragraph(subtitle.get_text(separator="\n", strip=True))
            set_rtl(p)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.color.rgb = RGBColor(128, 128, 128)
                run.italic = True
    
    doc.add_paragraph() # spacing
    
    if not doc_body:
        print("Could not find doc-body")
        return
        
    for child in doc_body.children:
        if child.name == 'h2':
            p = doc.add_heading(child.get_text(), level=2)
            set_rtl(p)
        elif child.name == 'p':
            p = doc.add_paragraph()
            set_rtl(p)
            for elem in child.children:
                if elem.name is None: # string
                    run = p.add_run(str(elem).replace('\n', ' '))
                elif elem.name == 'b' or elem.name == 'strong':
                    run = p.add_run(elem.get_text().replace('\n', ' '))
                    run.bold = True
                elif elem.name == 'span' and 'gold-mark' in elem.get('class', []):
                    run = p.add_run(elem.get_text().replace('\n', ' '))
                    run.font.color.rgb = RGBColor(146, 64, 14)
                    run.bold = True
                elif elem.name == 'a':
                    run = p.add_run(elem.get_text().replace('\n', ' '))
                    run.font.color.rgb = RGBColor(0, 0, 255)
                    run.underline = True
                else:
                    run = p.add_run(elem.get_text().replace('\n', ' '))
        elif child.name == 'ul':
            for li in child.find_all('li'):
                p = doc.add_paragraph(style='List Bullet')
                set_rtl(p)
                p.add_run(li.get_text().replace('\n', ' '))
        elif child.name == 'hr':
            p = doc.add_paragraph()
            run = p.add_run('----------------------------------------------------')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif child.name == 'div' and 'draft-notice' in child.get('class', []):
            p = doc.add_paragraph()
            set_rtl(p)
            run = p.add_run(child.get_text(separator=" ", strip=True))
            run.font.color.rgb = RGBColor(146, 64, 14)
            run.italic = True
        elif child.name == 'div' and 'doc-meta' in child.get('class', []):
            p = doc.add_paragraph()
            set_rtl(p)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(child.get_text(separator=" | ", strip=True))
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(128, 128, 128)

    doc.save(out_path)
    print(f"Saved {out_path}")

if __name__ == '__main__':
    html_file = r'c:\GitHub\tivuta\frontend\public\terms\index.html'
    out_file = r'c:\GitHub\tivuta\Tivuta_Terms_And_Conditions.docx'
    make_word_doc(html_file, out_file)
